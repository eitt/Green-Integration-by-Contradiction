from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from chapter_project.project_config import load_config


COMMENT_START = "<!--"
COMMENT_END = "-->"


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)
    image_path: str = ""


def clean_inline(text: str) -> str:
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1 (\2)", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"__(.*?)__", r"\1", text)
    text = re.sub(r"_(.*?)_", r"\1", text)
    return " ".join(text.split()).strip()


def is_table_separator(line: str) -> bool:
    trimmed = line.strip().strip("|").strip()
    return bool(trimmed) and all(part.strip("-: ") == "" for part in trimmed.split("|"))


def parse_table_row(line: str) -> list[str]:
    return [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]


def parse_markdown(markdown_text: str, base_dir: Path | None = None) -> list[Block]:
    lines = markdown_text.splitlines()
    blocks: list[Block] = []
    paragraph_lines: list[str] = []
    list_items: list[str] = []
    list_kind: str | None = None
    table_lines: list[str] = []
    in_comment = False

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if paragraph_lines:
            text = clean_inline(" ".join(paragraph_lines))
            if text:
                blocks.append(Block(kind="paragraph", text=text))
        paragraph_lines = []

    def flush_list() -> None:
        nonlocal list_items, list_kind
        if list_items and list_kind:
            blocks.append(Block(kind=list_kind, items=[clean_inline(item) for item in list_items]))
        list_items = []
        list_kind = None

    def flush_table() -> None:
        nonlocal table_lines
        if table_lines:
            rows = [parse_table_row(line) for line in table_lines if not is_table_separator(line)]
            if rows:
                blocks.append(Block(kind="table", rows=rows))
        table_lines = []

    for raw_line in lines:
        stripped = raw_line.strip()

        if in_comment:
            if COMMENT_END in stripped:
                in_comment = False
            continue

        if COMMENT_START in stripped:
            if COMMENT_END not in stripped:
                in_comment = True
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph()
            flush_list()
            flush_table()
            blocks.append(
                Block(
                    kind="heading",
                    level=len(heading_match.group(1)),
                    text=clean_inline(heading_match.group(2)),
                )
            )
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", stripped)
        if image_match:
            flush_paragraph()
            flush_list()
            flush_table()
            image_ref = image_match.group(2).strip()
            image_path = Path(image_ref)
            if base_dir is not None and not image_path.is_absolute():
                image_path = (base_dir / image_path).resolve()
            blocks.append(
                Block(
                    kind="image",
                    text=clean_inline(image_match.group(1)),
                    image_path=str(image_path),
                )
            )
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        number_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph()
            flush_table()
            if list_kind not in (None, "bullet_list"):
                flush_list()
            list_kind = "bullet_list"
            list_items.append(bullet_match.group(1))
            continue

        if number_match:
            flush_paragraph()
            flush_table()
            if list_kind not in (None, "number_list"):
                flush_list()
            list_kind = "number_list"
            list_items.append(number_match.group(1))
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_paragraph()
            flush_list()
            table_lines.append(stripped)
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            flush_table()
            continue

        flush_list()
        flush_table()
        paragraph_lines.append(stripped)

    flush_paragraph()
    flush_list()
    flush_table()
    return blocks


def set_style_font(style, font_name: str, size_pt: int, bold: bool = False, italic: bool = False) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic


def get_or_add_style(document: Document, style_name: str):
    try:
        return document.styles[style_name]
    except KeyError:
        return document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)


def configure_styles(document: Document, config: dict) -> None:
    call_format = config["call_format"]
    font_name = call_format["font_name"]
    font_size = call_format["font_size_pt"]

    normal = document.styles["Normal"]
    set_style_font(normal, font_name, font_size)
    normal.paragraph_format.line_spacing = call_format["line_spacing"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    title_style = get_or_add_style(document, "ChapterTitle")
    title_style.base_style = normal
    set_style_font(title_style, font_name, font_size, bold=True)
    title_style.paragraph_format.first_line_indent = Inches(0)
    title_style.paragraph_format.space_after = Pt(12)

    heading_1 = get_or_add_style(document, "ChapterHeading1")
    heading_1.base_style = normal
    set_style_font(heading_1, font_name, font_size, bold=True)
    heading_1.paragraph_format.first_line_indent = Inches(0)
    heading_1.paragraph_format.space_before = Pt(12)
    heading_1.paragraph_format.space_after = Pt(6)

    heading_2 = get_or_add_style(document, "ChapterHeading2")
    heading_2.base_style = normal
    set_style_font(heading_2, font_name, font_size, bold=True, italic=True)
    heading_2.paragraph_format.first_line_indent = Inches(0)
    heading_2.paragraph_format.space_before = Pt(6)
    heading_2.paragraph_format.space_after = Pt(3)

    heading_3 = get_or_add_style(document, "ChapterHeading3")
    heading_3.base_style = normal
    set_style_font(heading_3, font_name, font_size, bold=True)
    heading_3.paragraph_format.first_line_indent = Inches(0)
    heading_3.paragraph_format.space_before = Pt(6)
    heading_3.paragraph_format.space_after = Pt(3)

    no_indent = get_or_add_style(document, "BodyNoIndent")
    no_indent.base_style = normal
    set_style_font(no_indent, font_name, font_size)
    no_indent.paragraph_format.first_line_indent = Inches(0)

    reference_style = get_or_add_style(document, "ReferenceEntry")
    reference_style.base_style = normal
    set_style_font(reference_style, font_name, font_size)
    reference_style.paragraph_format.left_indent = Inches(0.5)
    reference_style.paragraph_format.first_line_indent = Inches(-0.5)


def set_run_font(run, font_name: str, size_pt: int) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, "Arial", 10)

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run._r.extend([begin, instr, separate, text, end])


def add_text_paragraph(document: Document, text: str, style_name: str, align: WD_ALIGN_PARAGRAPH):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_run_font(run, "Arial", 12)
    paragraph.alignment = align
    return paragraph


def add_heading(document: Document, block: Block) -> str:
    if block.level == 1:
        style_name = "ChapterTitle"
        alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif block.level == 2:
        style_name = "ChapterHeading1"
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif block.level == 3:
        style_name = "ChapterHeading2"
        alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        style_name = "ChapterHeading3"
        alignment = WD_ALIGN_PARAGRAPH.LEFT

    add_text_paragraph(document, block.text, style_name, alignment)
    return block.text.lower()


def add_list_block(document: Document, items: list[str], style_name: str) -> None:
    for item in items:
        paragraph = document.add_paragraph(style=style_name)
        run = paragraph.add_run(item)
        set_run_font(run, "Arial", 12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, "Arial", 12)
    run.font.bold = bold


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        padded = row + [""] * (column_count - len(row))
        for col_index, cell_text in enumerate(padded):
            set_cell_text(table.cell(row_index, col_index), cell_text, bold=row_index == 0)


def paragraph_style_for_section(section_name: str) -> tuple[str, WD_ALIGN_PARAGRAPH]:
    if section_name == "abstract":
        return "BodyNoIndent", WD_ALIGN_PARAGRAPH.JUSTIFY
    if section_name in {"keywords", "author profiles"}:
        return "BodyNoIndent", WD_ALIGN_PARAGRAPH.LEFT
    if section_name == "references":
        return "ReferenceEntry", WD_ALIGN_PARAGRAPH.LEFT
    return "Normal", WD_ALIGN_PARAGRAPH.JUSTIFY


def build_document(blocks: list[Block], config: dict) -> Document:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    configure_styles(document, config)
    add_page_number(section.footer.paragraphs[0])

    document.core_properties.title = config.get("suggested_title", "")
    document.core_properties.subject = config.get("project_name", "")

    current_section = ""
    for block in blocks:
        if block.kind == "heading":
            current_section = add_heading(document, block)
            continue

        if block.kind == "paragraph":
            style_name, alignment = paragraph_style_for_section(current_section)
            add_text_paragraph(document, block.text, style_name, alignment)
            continue

        if block.kind == "bullet_list":
            add_list_block(document, block.items, "List Bullet")
            continue

        if block.kind == "number_list":
            add_list_block(document, block.items, "List Number")
            continue

        if block.kind == "table":
            add_table(document, block.rows)
            continue

        if block.kind == "image":
            image_path = Path(block.image_path)
            if image_path.exists():
                paragraph = document.add_paragraph(style="BodyNoIndent")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=Inches(6.0))

    return document


def generate_docx(markdown_path: Path, config_path: Path, output_path: Path) -> Path:
    config = load_config(config_path)
    blocks = parse_markdown(markdown_path.read_text(encoding="utf-8"), base_dir=markdown_path.parent)
    document = build_document(blocks, config)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path.resolve()
