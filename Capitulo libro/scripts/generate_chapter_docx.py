from __future__ import annotations

import argparse
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from chapter_project.markdown_to_docx import generate_docx as generate_docx_python
from chapter_project.project_config import load_config


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate a formatted DOCX chapter document from the markdown draft."
    )
    parser.add_argument(
        "--source",
        default="draft/chapter_draft.md",
        help="Markdown source file.",
    )
    parser.add_argument(
        "--config",
        default="config/chapter_project.json",
        help="Project configuration file.",
    )
    parser.add_argument(
        "--output",
        default="output/doc/chapter_submission_template.docx",
        help="Output DOCX file.",
    )
    parser.add_argument(
        "--reference-doc",
        default="output/doc/reference_unab.docx",
        help="Reference DOCX used by pandoc.",
    )
    parser.add_argument(
        "--engine",
        choices=["auto", "pandoc", "python-docx"],
        default="auto",
        help="Conversion engine. Auto prefers pandoc when available.",
    )
    return parser.parse_args()


def set_style_font(style, font_name: str, size_pt: int, bold: bool = False, italic: bool = False) -> None:
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), font_name)
    r_fonts.set(qn("w:hAnsi"), font_name)
    r_fonts.set(qn("w:eastAsia"), font_name)
    style.font.name = font_name
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.italic = italic


def build_reference_docx(config_path: Path, reference_path: Path) -> Path:
    config = load_config(config_path)
    call_format = config["call_format"]
    font_name = call_format["font_name"]
    font_size = call_format["font_size_pt"]

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    set_style_font(normal, font_name, font_size)
    normal.paragraph_format.line_spacing = call_format["line_spacing"]
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.first_line_indent = Inches(0.5)

    heading1 = document.styles["Heading 1"]
    set_style_font(heading1, font_name, font_size, bold=True)
    heading1.paragraph_format.first_line_indent = Inches(0)
    heading1.paragraph_format.space_before = Pt(12)
    heading1.paragraph_format.space_after = Pt(6)

    heading2 = document.styles["Heading 2"]
    set_style_font(heading2, font_name, font_size, bold=True, italic=True)
    heading2.paragraph_format.first_line_indent = Inches(0)
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(3)

    heading3 = document.styles["Heading 3"]
    set_style_font(heading3, font_name, font_size, bold=True)
    heading3.paragraph_format.first_line_indent = Inches(0)
    heading3.paragraph_format.space_before = Pt(6)
    heading3.paragraph_format.space_after = Pt(3)

    title = document.styles["Title"]
    set_style_font(title, font_name, font_size, bold=True)
    title.paragraph_format.first_line_indent = Inches(0)

    document.add_paragraph("UNAB Chapter Reference", style="Title")
    document.add_paragraph("Reference styles for pandoc output.", style="Normal")

    reference_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(reference_path)
    return reference_path.resolve()


def generate_docx_pandoc(markdown_path: Path, output_path: Path, reference_path: Path) -> Path:
    markdown_path = markdown_path.resolve()
    output_path = output_path.resolve()
    reference_path = reference_path.resolve()
    cmd = [
        "pandoc",
        markdown_path.name,
        "--from",
        "gfm",
        "--to",
        "docx",
        "--standalone",
        "--reference-doc",
        str(reference_path),
        "--output",
        str(output_path),
    ]
    subprocess.run(cmd, check=True, cwd=markdown_path.parent)
    return output_path.resolve()


def main() -> int:
    args = parse_args()
    source_path = ROOT_DIR / args.source
    config_path = ROOT_DIR / args.config
    output_path = ROOT_DIR / args.output
    reference_path = ROOT_DIR / args.reference_doc

    engine = args.engine
    if engine == "auto":
        engine = "pandoc" if shutil.which("pandoc") else "python-docx"

    if engine == "pandoc":
        build_reference_docx(config_path, reference_path)
        output = generate_docx_pandoc(source_path, output_path, reference_path)
        print(f"DOCX generated with pandoc: {output}")
        return 0

    output = generate_docx_python(
        markdown_path=source_path,
        config_path=config_path,
        output_path=output_path,
    )
    print(f"DOCX generated with python-docx: {output}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
