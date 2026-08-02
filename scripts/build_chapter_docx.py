from pathlib import Path
import csv
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
TEXT = ROOT / 'draft' / 'chapter_text.md'
OUT = ROOT / 'draft' / 'GreenIntegration_RHandbook_V2.docx'


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_text(cell, text, bold=False, size=8.5):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Cambria'
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_field(paragraph, field):
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = field
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run('Page ')
    add_field(p, 'PAGE')


def setup_styles(doc):
    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Cambria'
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    title = styles['Title']
    title.font.name = 'Cambria'
    title.font.size = Pt(20)
    title.font.bold = True
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    subtitle = styles['Subtitle']
    subtitle.font.name = 'Cambria'
    subtitle.font.size = Pt(13)
    subtitle.font.italic = False
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)

    for name, size in [('Heading 1', 14), ('Heading 2', 12), ('Heading 3', 11)]:
        st = styles[name]
        st.font.name = 'Cambria'
        st.font.size = Pt(size)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(14 if name == 'Heading 1' else 10)
        st.paragraph_format.space_after = Pt(5)
        st.paragraph_format.keep_with_next = True

    if 'Caption' not in styles:
        styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    cap = styles['Caption']
    cap.font.name = 'Cambria'
    cap.font.size = Pt(10)
    cap.font.italic = True
    cap.paragraph_format.space_after = Pt(4)

    if 'Bibliography' not in styles:
        styles.add_style('Bibliography', WD_STYLE_TYPE.PARAGRAPH)
    bib = styles['Bibliography']
    bib.font.name = 'Cambria'
    bib.font.size = Pt(10.5)
    bib.paragraph_format.left_indent = Inches(0.25)
    bib.paragraph_format.first_line_indent = Inches(-0.25)
    bib.paragraph_format.space_after = Pt(4)


def add_body_paragraph(doc, text, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    return p


def add_timeline_figure(doc, caption):
    caption_paragraph = add_body_paragraph(doc, caption, 'Caption')
    caption_paragraph.paragraph_format.keep_with_next = True
    image_path = ROOT / 'figures' / 'EU-Mercosur_Negotiation_Timeline_polished.png'
    inline = doc.add_picture(str(image_path), width=Inches(6.55))
    inline._inline.docPr.set(
        'descr',
        'Black-and-white textured timeline of EU-Mercosur negotiations, '
        'environmental regulation and implementation from 1999 to 2026.',
    )
    image_paragraph = doc.paragraphs[-1]
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_together = True
    image_paragraph.paragraph_format.keep_with_next = True


def parse_markdown():
    lines = TEXT.read_text(encoding='utf-8').splitlines()
    title = lines[0].removeprefix('# ').strip()
    abstract = []
    keywords = ''
    body = []
    mode = None
    for line in lines[1:]:
        if line == '## Abstract':
            mode = 'abstract'
            continue
        if line == '## Keywords':
            mode = 'keywords'
            continue
        if line == '## Introduction':
            mode = 'body'
        if mode == 'abstract' and line.strip():
            abstract.append(line.strip())
        elif mode == 'keywords' and line.strip():
            keywords = line.strip()
        elif mode == 'body':
            body.append(line)
    return title, abstract, keywords, body


def clean_bib_value(value):
    value = value.strip().rstrip(',').strip()
    if value.startswith('{') and value.endswith('}'):
        value = value[1:-1]
    return value.replace('{', '').replace('}', '').replace('\\&', '&').replace('--', '-')


def format_authors(raw):
    people = [clean_bib_value(part) for part in raw.split(' and ')]
    names = []
    for index, person in enumerate(people):
        if ',' in person:
            family, given = [part.strip() for part in person.split(',', 1)]
            names.append(f'{family}, {given}' if index == 0 else f'{given} {family}')
        else:
            names.append(person)
    if len(names) == 1:
        return names[0]
    if len(names) == 2:
        return f'{names[0]}, and {names[1]}'
    return ', '.join(names[:-1]) + f', and {names[-1]}'


def parse_bibliography():
    entries = []
    current = None
    for line in (ROOT / 'references.bib').read_text(encoding='utf-8').splitlines():
        start = re.match(r'@(\w+)\{([^,]+),', line)
        if start:
            current = {'type': start.group(1).lower(), 'key': start.group(2)}
            continue
        if current is None:
            continue
        if line.strip() == '}':
            entries.append(current)
            current = None
            continue
        field = re.match(r'\s*(\w+)\s*=\s*(.+)', line)
        if field:
            current[field.group(1).lower()] = clean_bib_value(field.group(2))
    return entries


def format_reference(entry):
    author = format_authors(entry.get('author', 'Unknown author'))
    year = entry.get('year', 'n.d.')
    title = entry.get('title', 'Untitled')
    doi_url = f"https://doi.org/{entry['doi']}" if entry.get('doi') else entry.get('url', '')
    suffix = f' {doi_url}.' if doi_url else ''
    kind = entry['type']
    if kind == 'article':
        journal = entry.get('journal', entry.get('journaltitle', ''))
        volume = entry.get('volume', '')
        number = f" ({entry['number']})" if entry.get('number') else ''
        pages = f": {entry['pages']}" if entry.get('pages') else ''
        return f'{author}. {year}. "{title}." {journal} {volume}{number}{pages}.{suffix}'.replace('  ', ' ')
    if kind == 'book':
        place = entry.get('location', '')
        publisher = entry.get('publisher', '')
        publication = f'{place}: {publisher}' if place else publisher
        return f'{author}. {year}. {title}. {publication}.{suffix}'.replace('  ', ' ')
    if kind == 'incollection':
        pages = f", {entry['pages']}" if entry.get('pages') else ''
        place = entry.get('location', '')
        publisher = entry.get('publisher', '')
        publication = f'{place}: {publisher}' if place else publisher
        return f'{author}. {year}. "{title}." In {entry.get("booktitle", "")}{pages}. {publication}.{suffix}'.replace('  ', ' ')
    if kind == 'techreport':
        return f'{author}. {year}. {title}. {entry.get("institution", "")}.{suffix}'.replace('  ', ' ')
    if kind == 'legislation':
        journal = entry.get('journaltitle', '')
        pages = f", {entry['pages']}" if entry.get('pages') else ''
        return f'{author}. {year}. {title}. {journal}{pages}.{suffix}'.replace('  ', ' ')
    number = f" {entry['number']}." if entry.get('number') else ''
    return f'{author}. {year}. "{title}."{number}{suffix}'.replace('  ', ' ')


REFERENCES = [format_reference(entry) for entry in parse_bibliography()]


def build():
    doc = Document()
    setup_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.95)
    sec.right_margin = Inches(0.95)
    add_page_number(sec)

    doc.core_properties.title = 'Green Integration by Contradiction: EU Trade Policy and Latin America'
    doc.core_properties.subject = 'First draft for Routledge Handbook of European Integrations, Volume 2'
    doc.core_properties.author = 'Author to be confirmed'
    doc.core_properties.comments = 'Revised draft prepared 30 July 2026 from the chapter mind map, local source corpus and online reference checks.'
    doc.core_properties.keywords = 'European integration; EU-Mercosur; EU-Andean; sustainability governance; deforestation; trade policy'

    title, abstract, keywords, body = parse_markdown()
    p = doc.add_paragraph(style='Title')
    p.add_run(title)
    p = doc.add_paragraph(style='Subtitle')
    p.add_run('First draft for the Routledge Handbook of European Integrations, Volume 2')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Revised draft prepared 30 July 2026 | Author name(s) to be confirmed')
    r.font.size = Pt(10)
    r.italic = True

    doc.add_paragraph('Abstract', style='Heading 1')
    for para in abstract:
        add_body_paragraph(doc, para)
    doc.add_paragraph('Keywords', style='Heading 1')
    add_body_paragraph(doc, keywords)
    doc.add_page_break()

    for line in body:
        if not line.strip():
            continue
        if line.startswith('## '):
            add_body_paragraph(doc, line[3:].strip(), 'Heading 1')
        elif line.startswith('### '):
            add_body_paragraph(doc, line[4:].strip(), 'Heading 2')
        elif line.startswith('# '):
            continue
        elif match := re.fullmatch(r'!\[(.+)\]\((.+)\)', line.strip()):
            add_timeline_figure(doc, match.group(1))
        elif line.startswith('*Source:') and line.endswith('*'):
            add_body_paragraph(doc, line.strip('*'), 'Caption')
        else:
            add_body_paragraph(doc, line.strip())

    doc.add_page_break()
    add_body_paragraph(doc, 'References', 'Heading 1')
    for ref in REFERENCES:
        add_body_paragraph(doc, ref, 'Bibliography')

    doc.add_page_break()
    add_body_paragraph(doc, 'Table 1. Comparative sustainability provisions in the two agreements', 'Caption')
    table_rows = list(csv.DictReader((ROOT / 'audit' / 'table_10_provisions.csv').open(encoding='utf-8')))
    columns = ['dimension', 'EU-Andean agreement', 'EU-Mercosur agreement', 'interpretive relevance']
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [1.15, 2.0, 2.25, 1.3]
    for i, col in enumerate(columns):
        set_cell_text(table.rows[0].cells[i], col.replace('_', ' ').title(), bold=True, size=8.3)
        set_cell_shading(table.rows[0].cells[i], 'D9E2F3')
        table.rows[0].cells[i].width = Inches(widths[i])
    for row in table_rows:
        cells = table.add_row().cells
        for i, col in enumerate(columns):
            set_cell_text(cells[i], row[col], size=8.1)
            cells[i].width = Inches(widths[i])
    add_body_paragraph(doc, "Source: Authors' comparative synthesis based on European Commission (2023, 2026a, 2026c), Marín Durán (2020), Garcia and Gomez Arana (2025), Cesar de Oliveira et al. (2024), and Albertoni and Schlegelmilch (2026).", 'Caption')

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == '__main__':
    build()
