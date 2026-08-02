from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "draft" / "GreenIntegration_RHandbook_V2.docx"
FIGURE = ROOT / "figures" / "EU-Mercosur_Negotiation_Timeline_polished.png"


TABLE_ROWS = [
    (
        "Enforcement mechanisms",
        "Cooperative and managerial. Implementation relies on dialogue in the Trade "
        "and Sustainable Development Sub-committee, monitoring by Domestic Advisory "
        "Groups and recurring evaluation. The chapter does not provide for automatic "
        "trade sanctions (European Commission 2023; Garcia and Gomez Arana 2025).",
        "Strengthened but differentiated. The Additional Instrument clarifies "
        "implementation commitments, while the TSD Sub-committee, Domestic Advisory "
        "Groups, consultations and a panel of experts remain central. The Paris "
        "Agreement has a separate status as an essential element (European Commission "
        "2026c).",
    ),
    (
        "Dispute resolution",
        "Government consultations may be followed by a panel of experts. Its public "
        "report contains findings and recommendations, but the TSD chapter is excluded "
        "from the agreement's general trade dispute-settlement route.",
        "Government consultations may be followed by an independent panel of experts. "
        "Trade sanctions do not follow automatically from every TSD breach. Suspension "
        "is linked instead to a serious breach of, or withdrawal from, the Paris "
        "Agreement through the essential-elements clause (European Commission 2026c; "
        "Akdogan 2025).",
    ),
    (
        "Conditionality design",
        "Cooperative externalisation. Environmental and labour objectives are framed "
        "as shared commitments and pursued through dialogue, monitoring and technical "
        "cooperation. Continued preferences are not routinely conditioned on measured "
        "sustainability outcomes.",
        "Integration by contradiction. Market opening is combined with stronger "
        "deforestation commitments and an essential-elements climate clause. The "
        "legal signal is firmer, but implementation still depends on domestic "
        "institutions, evidence and financial capacity.",
    ),
    (
        "Sectoral coverage",
        "Broad TSD coverage includes labour standards, multilateral environmental "
        "agreements, biodiversity and the sustainable use of natural resources. "
        "Sector-specific concerns are developed mainly through committee work and "
        "evaluation.",
        "The agreement gives greater political weight to deforestation, land use, "
        "agriculture and biodiversity, while also promoting sustainable trade and "
        "investment in critical raw materials. The EUDR remains a parallel EU "
        "regulation rather than part of the treaty.",
    ),
]


def paragraph_after(document, anchor, style=None):
    paragraph = document.add_paragraph(style=style)
    anchor._p.addnext(paragraph._p)
    return paragraph


def element_after(anchor_element, new_element):
    anchor_element.addnext(new_element)
    return new_element


def add_field(paragraph, instruction, result):
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instruction_run = paragraph.add_run()
    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction
    instruction_run._r.append(instruction_text)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    result_run = paragraph.add_run(result)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    return result_run


def add_reference(paragraph, bookmark, display):
    add_field(paragraph, f" REF {bookmark} \\h ", display)


def add_numbered_caption(paragraph, label, bookmark, title, bookmark_id):
    paragraph.clear()
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark)
    paragraph._p.append(start)

    paragraph.add_run(f"{label} ")
    add_field(paragraph, f" SEQ {label} \\* ARABIC ", "1")

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.append(end)
    paragraph.add_run(f". {title}")
    paragraph.paragraph_format.keep_with_next = True


def add_update_fields_setting(document):
    settings = document.settings._element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def remove_existing_end_table(document):
    for paragraph in list(document.paragraphs):
        if (
            paragraph.text.startswith("Table 1. Comparative sustainability provisions")
            or paragraph.text.startswith(
                "Source: Authors' comparative synthesis based on European Commission"
            )
        ):
            paragraph._element.getparent().remove(paragraph._element)
    for table in list(document.tables):
        table._element.getparent().remove(table._element)


def set_cell_text(cell, text, bold=False, white=False, size=8.4):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Cambria"
    run.font.size = Pt(size)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_header(row):
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_cant_split(row):
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def find_paragraph(document, exact_text):
    for paragraph in document.paragraphs:
        if paragraph.text == exact_text:
            return paragraph
    raise RuntimeError(f"Paragraph not found: {exact_text}")


def main():
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)

    document = Document(DOCX)
    if document._element.xpath('.//w:bookmarkStart[@w:name="_RefFigure1"]'):
        raise RuntimeError("Figure and table cross-references are already present.")

    figure_anchor = find_paragraph(
        document,
        "Implementation remains too recent to support an ex post assessment of "
        "environmental effects. The legal sequence can be documented, but claims "
        "about ecological transformation would be premature. The analysis must "
        "therefore examine the obligations, procedures and expectations created by "
        "the revised architecture while separating observed institutional "
        "developments from projected outcomes.",
    )
    table_anchor = find_paragraph(
        document,
        "Mercosur reveals the more openly political form. Stronger language on "
        "climate and deforestation became necessary, in part, because the agreement "
        "faced a legitimacy crisis. Environmental provisions helped achieve political "
        "closure by addressing criticism and allowing the parties to present market "
        "opening as compatible with climate commitments. Those provisions may have "
        "substantive effects, but their inclusion does not resolve the agricultural "
        "and land-use conflicts embedded in the commercial relationship. The "
        "comparison therefore shifts attention from whether sustainability language "
        "exists to what institutional and political work it performs.",
    )

    remove_existing_end_table(document)

    figure_reference = paragraph_after(document, figure_anchor)
    add_reference(figure_reference, "_RefFigure1", "Figure 1")
    figure_reference.add_run(
        " places the legal sequence alongside the Green Deal and the EUDR. It is a "
        "chronological guide to the overlapping instruments, not evidence of their "
        "environmental effects."
    )

    figure_caption = paragraph_after(document, figure_reference, "Caption")
    add_numbered_caption(
        figure_caption,
        "Figure",
        "_RefFigure1",
        "EU-Mercosur: negotiation, green commitments and implementation, 1999-2026",
        101,
    )
    figure_caption.paragraph_format.page_break_before = True

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_together = True
    image_paragraph.paragraph_format.keep_with_next = True
    image = image_paragraph.add_run().add_picture(str(FIGURE), width=Inches(6.45))
    image._inline.docPr.set(
        "descr",
        "Black-and-white textured timeline of EU-Mercosur negotiations, "
        "environmental regulation and implementation from 1999 to 2026.",
    )
    element_after(figure_caption._p, image_paragraph._p)

    figure_source = document.add_paragraph(style="Caption")
    figure_source.add_run(
        "Source: Authors' compilation from the European Parliament's negotiation "
        "history, European Commission agreement texts and chronology, Regulation "
        "(EU) 2023/1115, Albertoni and Schlegelmilch (2026), and Cesar de Oliveira "
        "et al. (2024)."
    )
    element_after(image_paragraph._p, figure_source._p)

    table_reference = paragraph_after(document, table_anchor)
    add_reference(table_reference, "_RefTable1", "Table 1")
    table_reference.add_run(
        " compares the institutional design of the two agreements before the "
        "analysis turns to asymmetric greening. It distinguishes ordinary TSD "
        "procedures from the separate essential-elements clause in the Mercosur "
        "agreement."
    )

    table_caption = paragraph_after(document, table_reference, "Caption")
    add_numbered_caption(
        table_caption,
        "Table",
        "_RefTable1",
        "Comparative sustainability provisions in the EU-Andean and EU-Mercosur agreements",
        102,
    )
    table_caption.paragraph_format.page_break_before = True

    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.25), Inches(2.58), Inches(2.58)]
    headers = [
        "Feature",
        "EU-Andean Trade Agreement",
        "EU-Mercosur Agreement (2024 outcome)",
    ]
    for index, text in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_text(cell, text, bold=True, white=True, size=8.5)
        shade_cell(cell, "333333")
        cell.width = widths[index]
    set_repeat_header(table.rows[0])
    set_cant_split(table.rows[0])

    for feature, andean, mercosur in TABLE_ROWS:
        cells = table.add_row().cells
        values = [feature, andean, mercosur]
        for index, text in enumerate(values):
            set_cell_text(cells[index], text, bold=index == 0, size=8.2)
            cells[index].width = widths[index]
            if index == 0:
                shade_cell(cells[index], "E6E6E6")
        set_cant_split(table.rows[-1])

    element_after(table_caption._p, table._element)

    table_source = document.add_paragraph(style="Caption")
    table_source.add_run(
        "Source: Authors' synthesis based on European Commission (2023, 2026a, "
        "2026c), Garcia and Gomez Arana (2025), Akdogan (2025), Cesar de Oliveira "
        "et al. (2024), and the final agreement text."
    )
    element_after(table._element, table_source._p)

    add_update_fields_setting(document)
    document.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
