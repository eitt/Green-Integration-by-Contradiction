from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "draft" / "GreenIntegration_RHandbook_V2.docx"
BACKUP_PATH = (
    ROOT
    / "draft"
    / "GreenIntegration_RHandbook_V2_before_followup_audit_20260730.docx"
)


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def replace_paragraph(document: Document, prefix: str, old: str, new: str) -> None:
    paragraph = find_paragraph(document, prefix)
    if old in paragraph.text:
        set_paragraph_text(paragraph, paragraph.text.replace(old, new))
    elif new not in paragraph.text:
        raise RuntimeError(f"Expected text not found in paragraph: {prefix}")


def main() -> None:
    if not BACKUP_PATH.exists():
        copy2(DOCX_PATH, BACKUP_PATH)

    document = Document(DOCX_PATH)

    replace_paragraph(
        document,
        "The central argument is deliberately counterintuitive",
        (
            "In some circumstances, sustainability language operates primarily as "
            "a legitimacy device: it helps an agreement appear compatible with "
            "European values and contemporary climate politics while leaving the "
            "distribution of adjustment costs unresolved (Garcia & Gomez Arana, "
            "2025; Judijanto, 2025; Marín Durán, 2020). In other circumstances"
        ),
        (
            "In some circumstances, sustainability language operates primarily as "
            "a legitimacy device: it helps an agreement appear compatible with "
            "European values and contemporary climate politics while leaving the "
            "distribution of adjustment costs unresolved (Garcia & Gomez Arana, "
            "2025; Marín Durán, 2020). Evidence from the EUDR shows a related "
            "distributive problem: compliance burdens may fall unevenly on "
            "smallholders in the Global South (Judijanto, 2025). In other circumstances"
        ),
    )

    replace_paragraph(
        document,
        "The EU Deforestation Regulation provides parallel evidence",
        (
            "on the basis of evidence that they are not associated with recent "
            "deforestation or forest degradation."
        ),
        (
            "on the basis of evidence that they are not associated with recent "
            "deforestation or forest degradation (European Parliament and Council "
            "2023; European Commission 2026b)."
        ),
    )

    replace_paragraph(
        document,
        "Changes to the implementation timetable and information system",
        (
            "simplified procedures for certain small primary operators."
        ),
        (
            "simplified procedures for certain small primary operators (European "
            "Commission 2026b; European Commission 2026d)."
        ),
    )

    replacements = {
        "Judijanto, L. (2025).": (
            'Judijanto, Loso. 2025. "Green Neo-colonialism or Sustainable Trade: '
            "EUDR's Equity Implications for Global South Smallholders.\" "
            "International Journal of Environmental Sciences 11 (5): 994-1003."
        ),
        "Olech, I., Krupska, K., & Kosior, K. (2025).": (
            'Olech, Igor, Katarzyna Krupska, and Katarzyna Kosior. 2025. "The '
            "Problem of Enforcing Environmental Clauses in the EU-Mercosur "
            "Partnership Agreement in the Context of the Discrepancy in "
            'Deforestation Indices." Forests 16 (12): 1821. '
            "https://doi.org/10.3390/f16121821."
        ),
        "Von der Leyen, U. et al. (2020, 31 de enero).": (
            'Von der Leyen, Ursula, et al. 2020. "Joint Statement by the '
            "Presidents of the European Commission, the European Parliament and "
            'the European Council." 31 January 2020.'
        ),
    }

    for prefix, replacement in replacements.items():
        paragraph = find_paragraph(document, prefix)
        set_paragraph_text(paragraph, replacement)

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
