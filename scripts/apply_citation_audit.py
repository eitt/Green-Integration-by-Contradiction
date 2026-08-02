from __future__ import annotations

from pathlib import Path
from shutil import copy2

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "draft" / "GreenIntegration_RHandbook_V2.docx"
BACKUP_PATH = (
    ROOT
    / "draft"
    / "GreenIntegration_RHandbook_V2_before_citation_audit_20260730.docx"
)


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.startswith(prefix):
            return paragraph
    raise RuntimeError(f"Paragraph not found: {prefix}")


def replace_in_runs(paragraph: Paragraph, old: str, new: str) -> None:
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return
    raise RuntimeError(f"Text not found in runs: {old}")


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    if not BACKUP_PATH.exists():
        copy2(DOCX_PATH, BACKUP_PATH)

    document = Document(DOCX_PATH)

    introduction = find_paragraph(
        document, "European integration is often narrated as an internal project"
    )
    if "; IRIS, 2020" in introduction.text:
        replace_in_runs(introduction, "; IRIS, 2020", "")

    argument = find_paragraph(
        document, "The central argument is deliberately counterintuitive"
    )
    old_tail = (
        "The difference cannot be read from the presence of a sustainability "
        "chapter alone, because the similarity between agreement texts with "
        "different policies is about 80%. (Cesar de Oliveira et al., 2024)."
    )
    new_tail = (
        "The difference cannot be inferred from the presence of a sustainability "
        "chapter alone: a shared institutional vocabulary can perform different "
        "political work under different implementation conditions."
    )
    full_text = argument.text
    if old_tail in full_text:
        set_paragraph_text(argument, full_text.replace(old_tail, new_tail))
    elif new_tail not in full_text:
        raise RuntimeError("Neither the unsupported claim nor its correction was found")

    bradford_2020 = find_paragraph(document, "Bradford, Anu. 2020.")
    if not any(p.text.startswith("Bradford, Anu. 2012.") for p in document.paragraphs):
        bradford_2020.insert_paragraph_before(
            'Bradford, Anu. 2012. "The Brussels Effect." Northwestern University '
            "Law Review 107 (1): 1-68. "
            "https://scholarlycommons.law.northwestern.edu/nulr/vol107/iss1/1.",
            style=bradford_2020.style,
        )

    lavenex_2009 = find_paragraph(
        document, "Lavenex, Sandra, and Frank Schimmelfennig. 2009."
    )
    if not any(p.text.startswith("Lavenex, Sandra. 2004.") for p in document.paragraphs):
        lavenex_2009.insert_paragraph_before(
            'Lavenex, Sandra. 2004. "EU External Governance in Wider Europe." '
            "Journal of European Public Policy 11 (4): 680-700. "
            "https://doi.org/10.1080/1350176042000248098.",
            style=lavenex_2009.style,
        )

    def find_ec(suffix: str, title_prefix: str) -> Paragraph:
        try:
            return find_paragraph(
                document, f'European Commission. 2026. "{title_prefix}'
            )
        except RuntimeError:
            return find_paragraph(
                document, f'European Commission. 2026{suffix}. "{title_prefix}'
            )

    ec_a = find_ec("a", "EU-Mercosur: Text of the Agreement.")
    ec_b = find_ec("b", "Regulation on Deforestation-free Products.")
    ec_d = find_ec(
        "d", "Commission Updates Product Scope and Digital Tools"
    )
    ec_c = find_ec(
        "c", "EU-Mercosur Partnership Agreement: Trade and Sustainable Development."
    )

    if ec_a.text.startswith("European Commission. 2026."):
        replace_in_runs(ec_a, "European Commission. 2026.", "European Commission. 2026a.")
    if ec_b.text.startswith("European Commission. 2026."):
        replace_in_runs(ec_b, "European Commission. 2026.", "European Commission. 2026b.")
    if ec_c.text.startswith("European Commission. 2026."):
        replace_in_runs(ec_c, "European Commission. 2026.", "European Commission. 2026c.")
    if ec_d.text.startswith("European Commission. 2026."):
        replace_in_runs(ec_d, "European Commission. 2026.", "European Commission. 2026d.")

    # Keep the manual reference list in the same a-b-c-d sequence used in the text.
    ec_d._element.addprevious(ec_c._element)

    for paragraph in list(document.paragraphs):
        if paragraph.text.startswith(
            "Regulation (EU) 2023/1115 of the European Parliament and of the Council"
        ):
            delete_paragraph(paragraph)

    document.save(DOCX_PATH)


if __name__ == "__main__":
    main()
